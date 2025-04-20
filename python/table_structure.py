import docx
import pandas as pd
import re
import time

def extract_course_info_with_modules_labs(file_path):
    """
    .docx 파일에서 과정 정보, 모듈, 실습명을 추출하는 함수
    
    Args:
        file_path (str): .docx 파일 경로
        
    Returns:
        dict: 추출된 과정 정보, 모듈, 실습명을 담은 딕셔너리
    """
    print(f"문서 제목: {file_path}")
    
    # 문서 열기
    doc = docx.Document(file_path)
    print(f"문서를 불러왔습니다.")
    
    # 기본 정보 출력
    print(f"문단 수: {len(doc.paragraphs)}")
    print(f"표 수: {len(doc.tables)}")
    
    # 결과를 저장할 딕셔너리
    result = {
        "two_row_tables": [],
        "course_modules": {},
        "course_labs": {}
    }
    
    # 1. 2줄짜리 표에서 레벨과 소요 시간 추출
    print("\n### 2줄짜리 표에서 레벨과 소요 시간 추출 중 ###")
    for idx, table in enumerate(doc.tables):
        try:
            # 진행 상황 로깅
            if idx % 20 == 0 and idx > 0:
                print(f"표 {idx}/{len(doc.tables)} 처리 중...")
            
            # 2줄짜리 표만 필터링
            if len(table.rows) == 2:
                print(f"\n2줄짜리 표 발견 (표 #{idx+1})")
                
                # 표 데이터 추출
                headers = []
                values = []
                
                # 첫 번째 행 (헤더)
                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip())
                
                # 두 번째 행 (값)
                for cell in table.rows[1].cells:
                    values.append(cell.text.strip())
                
                # 레벨과 소요 시간 정보 찾기
                level = None
                duration = None
                
                # 헤더와 값을 순회하며 찾기
                for i, header in enumerate(headers):
                    if i < len(values):  # 인덱스 범위 확인
                        header_lower = header.lower()
                        if "레벨" in header_lower or "수준" in header_lower:
                            level = values[i]
                        elif "소요 시간" in header_lower or "기간" in header_lower or "duration" in header_lower:
                            duration = values[i]
                
                # 과정명 추출 (표 주변 문단에서)
                course_name = extract_course_name_near_table(doc, table, idx, max_distance=10)
                
                # 결과 저장
                table_info = {
                    "table_index": idx,
                    "course_name": course_name,
                    "level": level,
                    "duration": duration,
                    "headers": headers,
                    "values": values
                }
                
                result["two_row_tables"].append(table_info)
                print(f"  추출 정보: 과정명='{course_name}', 레벨='{level}', 소요 시간='{duration}'")
                
        except Exception as e:
            print(f"표 {idx+1} 처리 중 오류 발생: {str(e)}")
    
    # 2. 모듈명 추출
    print("\n### 모듈명 추출 중 ###")
    course_modules = extract_modules(doc)
    result["course_modules"] = course_modules
    
    # 3. 실습명 추출
    print("\n### 실습명 추출 중 ###")
    course_labs = extract_labs(doc)
    result["course_labs"] = course_labs
    
    return result

def extract_course_name_near_table(doc, table, table_idx, max_distance=10):
    """
    표 주변 문단에서 과정명을 추출
    
    Args:
        doc: 문서 객체
        table: 표 객체
        table_idx: 표 인덱스
        max_distance: 표로부터 검색할 최대 문단 거리
        
    Returns:
        str: 추출된 과정명
    """
    # 표의 위치 찾기
    table_index_in_doc = -1
    doc_elements = []
    
    # 문서 요소 순회
    for para_idx, para in enumerate(doc.paragraphs):
        doc_elements.append(("para", para_idx, para))
    
    for tbl_idx, tbl in enumerate(doc.tables):
        doc_elements.append(("table", tbl_idx, tbl))
        if tbl_idx == table_idx:
            table_index_in_doc = len(doc_elements) - 1
    
    if table_index_in_doc == -1:
        return "과정명 추출 실패"
    
    # 표 전후로 문단 검색
    course_name_candidates = []
    
    # 표 이전 문단 확인
    start_idx = max(0, table_index_in_doc - max_distance)
    for i in range(table_index_in_doc-1, start_idx-1, -1):
        if i < 0 or i >= len(doc_elements) or doc_elements[i][0] != "para":
            continue
        
        para = doc_elements[i][2]
        text = para.text.strip()
        
        # 과정명으로 볼 수 있는 조건
        if text and 5 < len(text) < 150 and not text.startswith("모듈") and not text.startswith("실습"):
            if "AWS" in text or "Amazon" in text or "Cloud" in text:
                course_name_candidates.append(text)
            elif re.match(r'^[\w\s\-]+\$', text) and len(text.split()) <= 10:  # 일반 텍스트이며 단어 수가 적은 경우
                course_name_candidates.append(text)
    
    # 가장 적합한 후보 선택
    if course_name_candidates:
        return course_name_candidates[0]  # 표에 가장 가까운 후보 선택
    
    return "과정명 추출 실패"

def extract_modules(doc):
    """
    문서에서 모듈 정보 추출
    
    Args:
        doc: 문서 객체
        
    Returns:
        dict: 과정별 모듈 정보
    """
    modules = {}
    current_course = None
    current_modules = []
    
    # 모듈 패턴 (모듈 1: 제목 또는 Module 1: 제목)
    module_pattern = re.compile(r'모듈\s*(\d+)[:\s-]\s*(.+)', re.IGNORECASE)
    module_pattern_eng = re.compile(r'Module\s*(\d+)[:\s-]\s*(.+)', re.IGNORECASE)
    
    # 과정명 패턴 (특정 키워드를 포함하는 짧은 문단)
    course_pattern = re.compile(r'.*\b(AWS|Amazon|Cloud)\b.*', re.IGNORECASE)
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 과정명으로 보이는 문단 확인
        if (course_pattern.match(text) and 
            10 < len(text) < 100 and 
            not text.startswith("모듈") and 
            not text.startswith("실습") and
            not text.startswith("•")):
            
            # 이전 과정의 모듈을 저장
            if current_course and current_modules:
                modules[current_course] = current_modules
            
            # 새 과정 시작
            current_course = text
            current_modules = []
            continue
        
        # 모듈 패턴 확인
        module_match = module_pattern.match(text) or module_pattern_eng.match(text)
        if module_match and current_course:
            module_num = module_match.group(1)
            module_title = module_match.group(2).strip()
            
            # 모듈명이 너무 길거나 불릿 포인트로 시작하면 제외
            if len(module_title) < 200 and not module_title.startswith('•'):
                current_modules.append({
                    "module_number": module_num,
                    "module_title": module_title
                })
    
    # 마지막 과정의 모듈 저장
    if current_course and current_modules:
        modules[current_course] = current_modules
    
    return modules

def extract_labs(doc):
    """
    문서에서 실습 정보 추출
    
    Args:
        doc: 문서 객체
        
    Returns:
        dict: 과정별 실습 정보
    """
    labs = {}
    current_course = None
    current_labs = []
    
    # 실습 패턴 (실습 1: 제목 또는 Lab 1: 제목)
    lab_pattern = re.compile(r'실습\s*(\d+)[:\s-]\s*(.+)', re.IGNORECASE)
    lab_pattern_eng = re.compile(r'Lab\s*(\d+)[:\s-]\s*(.+)', re.IGNORECASE)
    
    # 과정명 패턴 (특정 키워드를 포함하는 짧은 문단)
    course_pattern = re.compile(r'.*\b(AWS|Amazon|Cloud)\b.*', re.IGNORECASE)
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 과정명으로 보이는 문단 확인
        if (course_pattern.match(text) and 
            10 < len(text) < 100 and 
            not text.startswith("모듈") and 
            not text.startswith("실습") and
            not text.startswith("•")):
            
            # 이전 과정의 실습을 저장
            if current_course and current_labs:
                labs[current_course] = current_labs
            
            # 새 과정 시작
            current_course = text
            current_labs = []
            continue
        
        # 실습 패턴 확인
        lab_match = lab_pattern.match(text) or lab_pattern_eng.match(text)
        if lab_match and current_course:
            lab_num = lab_match.group(1)
            lab_title = lab_match.group(2).strip()
            
            # 실습명이 너무 길거나 불릿 포인트로 시작하면 제외
            if len(lab_title) < 200 and not lab_title.startswith('•'):
                current_labs.append({
                    "lab_number": lab_num,
                    "lab_title": lab_title
                })
    
    # 마지막 과정의 실습 저장
    if current_course and current_labs:
        labs[current_course] = current_labs
    
    return labs

def display_course_info_preview(result):
    """
    추출된 과정 정보를 미리 보기로 출력
    
    Args:
        result: 추출된 과정 정보
    """
    two_row_tables = result["two_row_tables"]
    course_modules = result["course_modules"]
    course_labs = result["course_labs"]
    
    # 과정 정보 요약 DataFrame
    if two_row_tables:
        summary_data = []
        for info in two_row_tables:
            summary_data.append({
                "과정명": info["course_name"],
                "레벨": info["level"],
                "소요 시간": info["duration"]
            })
        
        df_summary = pd.DataFrame(summary_data)
        print("\n=== 과정 정보 요약 ===")
        print(df_summary)
        
        # CSV로 저장
        df_summary.to_csv("course_summary.csv", index=False, encoding='utf-8-sig')
        print("과정 정보 요약을 'course_summary.csv'에 저장했습니다.")
    
    # 각 과정별 모듈 및 실습 정보 출력
    print("\n=== 과정별 모듈 및 실습 정보 ===")
    all_courses = set()
    
    # 두 줄 표에서 추출한 과정명
    for info in two_row_tables:
        if info["course_name"] != "과정명 추출 실패":
            all_courses.add(info["course_name"])
    
    # 모듈에서 추출한 과정명
    for course in course_modules:
        all_courses.add(course)
    
    # 실습에서 추출한 과정명
    for course in course_labs:
        all_courses.add(course)
    
    # 모든 과정에 대한 정보 출력
    all_course_info = []
    
    for course in all_courses:
        course_info = {
            "과정명": course,
            "레벨": "",
            "소요 시간": "",
            "모듈 수": 0,
            "실습 수": 0
        }
        
        # 레벨과 소요 시간 추가
        for info in two_row_tables:
            if info["course_name"] == course:
                course_info["레벨"] = info["level"] if info["level"] else ""
                course_info["소요 시간"] = info["duration"] if info["duration"] else ""
                break
        
        # 모듈 정보 추가
        if course in course_modules:
            course_info["모듈 수"] = len(course_modules[course])
        
        # 실습 정보 추가
        if course in course_labs:
            course_info["실습 수"] = len(course_labs[course])
        
        all_course_info.append(course_info)
    
    # 과정별 종합 정보 DataFrame
    df_all = pd.DataFrame(all_course_info)
    print(df_all)
    
    # CSV로 저장
    df_all.to_csv("all_course_info.csv", index=False, encoding='utf-8-sig')
    print("모든 과정 정보를 'all_course_info.csv'에 저장했습니다.")
    
    # 특정 과정의 모듈 및 실습 정보 상세 출력 (선택적)
    if all_courses:
        print("\n=== 첫 번째 과정의 모듈 및 실습 상세 정보 ===")
        sample_course = list(all_courses)[0]
        print(f"과정명: {sample_course}")
        
        # 모듈 정보
        if sample_course in course_modules and course_modules[sample_course]:
            print("\n모듈 정보:")
            for module in course_modules[sample_course][:5]:  # 처음 5개만
                print(f"  모듈 {module['module_number']}: {module['module_title']}")
            if len(course_modules[sample_course]) > 5:
                print(f"  ... 외 {len(course_modules[sample_course]) - 5}개 모듈")
        
        # 실습 정보
        if sample_course in course_labs and course_labs[sample_course]:
            print("\n실습 정보:")
            for lab in course_labs[sample_course][:5]:  # 처음 5개만
                print(f"  실습 {lab['lab_number']}: {lab['lab_title']}")
            if len(course_labs[sample_course]) > 5:
                print(f"  ... 외 {len(course_labs[sample_course]) - 5}개 실습")
    
    # 모듈 및 실습 정보를 개별 파일로 저장
    modules_data = []
    for course, modules in course_modules.items():
        for module in modules:
            modules_data.append({
                "과정명": course,
                "모듈 번호": module["module_number"],
                "모듈 제목": module["module_title"]
            })
    
    if modules_data:
        pd.DataFrame(modules_data).to_csv("modules_info.csv", index=False, encoding='utf-8-sig')
        print("\n모듈 정보를 'modules_info.csv'에 저장했습니다.")
    
    labs_data = []
    for course, labs in course_labs.items():
        for lab in labs:
            labs_data.append({
                "과정명": course,
                "실습 번호": lab["lab_number"],
                "실습 제목": lab["lab_title"]
            })
    
    if labs_data:
        pd.DataFrame(labs_data).to_csv("labs_info.csv", index=False, encoding='utf-8-sig')
        print("실습 정보를 'labs_info.csv'에 저장했습니다.")

def main():
    """메인 실행 함수"""
    try:
        import sys
        if len(sys.argv) > 1:
            file_path = sys.argv[1]
        else:
            file_path = 'AWS TnC_ILT_DILT.docx'  # 기본 파일 이름
        
        start_time = time.time()
        result = extract_course_info_with_modules_labs(file_path)
        end_time = time.time()
        
        print(f"\n처리 완료! 소요 시간: {end_time - start_time:.2f}초")
        display_course_info_preview(result)
        
        # 문서 내용 전체를 저장 (선택적)
        doc = docx.Document(file_path)
        with open("extracted_text.txt", "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
            print("\n문서 내용을 extracted_text.txt 파일로 저장했습니다.")
            
    except Exception as e:
        print(f"오류 발생: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()