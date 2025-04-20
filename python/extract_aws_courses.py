import docx
import pandas as pd
import re
import time
import json
from collections import defaultdict
import boto3
import json
import uuid
from datetime import datetime
import time
from boto3.dynamodb.conditions import Key, Attr
from decimal import Decimal

def save_courses_to_dynamodb(json_file_path, region_name='ap-northeast-2'):
    """
    JSON 파일에서 과정 정보를 읽어 DynamoDB에 저장하는 함수
    
    Args:
        json_file_path (str): 과정 정보가 담긴 JSON 파일 경로
        region_name (str): AWS 리전 이름
    """
    # 데이터 로드
    try:
        with open(json_file_path, 'r', encoding='utf-8') as f:
            course_data = json.load(f)
        print(f"총 {len(course_data)}개 과정 정보를 로드했습니다.")
    except Exception as e:
        print(f"JSON 파일 로드 중 오류 발생: {str(e)}")
        return False
    
    # AWS 세션 및 DynamoDB 클라이언트 생성
    try:
        session = boto3.Session(region_name=region_name)
        dynamodb = session.resource('dynamodb')
        
        # 테이블 참조
        catalog_table = dynamodb.Table('Tnc-CourseCatalog')
        modules_table = dynamodb.Table('Tnc-CourseCatalog-Modules')
        
        print("AWS DynamoDB에 연결했습니다.")
    except Exception as e:
        print(f"AWS 연결 중 오류 발생: {str(e)}")
        return False
    
    # 저장 전 데이터 미리보기
    sample_course = course_data[0] if course_data else None
    if sample_course:
        print("\n=== 저장할 데이터 미리보기 ===")
        print(f"과정명: {sample_course['course_name']}")
        print(f"레벨: {sample_course['level']}")
        print(f"소요 시간: {sample_course['duration']}")
        print(f"모듈 수: {len(sample_course.get('modules', []))}")
        print(f"실습 수: {len(sample_course.get('labs', []))}")
    
    # 사용자 확인
    confirmation = input("\n위 정보를 DynamoDB에 저장하시겠습니까? (y/n): ")
    if confirmation.lower() != 'y':
        print("저장이 취소되었습니다.")
        return False
    
    # 타임스탬프 및 성공/실패 카운터
    timestamp = datetime.now().isoformat()
    success_count = 0
    failure_count = 0
    
    # 과정 정보를 DynamoDB 테이블에 저장
    print("\n=== 과정 정보 저장 중 ===")
    for course in course_data:
        try:
            # 과정 ID 생성 (고유 식별자)
            course_id = str(uuid.uuid4())
            
            # 과정 정보 항목 생성
            catalog_item = {
                "courseId": course_id,
                "courseName": course['course_name'],
                "level": course['level'],
                "duration": course['duration'],
                "deliveryMethod": course.get('delivery_method', ''),
                "description": course.get('description', ''),
                "objectives": course.get('objectives', ''),
                "audience": course.get('audience', ''),
                "prerequisites": course.get('prerequisites', ''),
                "outline": course.get('outline', ''),
                "createdAt": timestamp,
                "updatedAt": timestamp
            }
            
            # DynamoDB에 과정 정보 저장
            catalog_table.put_item(Item=catalog_item)
            
            print(f"과정 정보 저장 완료: {course['course_name']}")
            success_count += 1
            
            # 모듈 정보 저장
            module_count = 0
            for idx, module in enumerate(course.get('modules', [])):
                try:
                    module_id = f"{course_id}#module#{idx+1}"
                    module_item = {
                        "courseId": course_id,
                        "id": module_id,
                        "courseName": course['course_name'],
                        "moduleNumber": module['module_number'],
                        "moduleTitle": module['module_title'],
                        "type": "module",
                        "createdAt": timestamp,
                        "updatedAt": timestamp
                    }
                    
                    modules_table.put_item(Item=module_item)
                    module_count += 1
                except Exception as e:
                    print(f"모듈 저장 중 오류 발생: {str(e)}")
                    failure_count += 1
            
            # 실습 정보 저장
            lab_count = 0
            for idx, lab in enumerate(course.get('labs', [])):
                try:
                    lab_id = f"{course_id}#lab#{idx+1}"
                    lab_item = {
                        "courseId": course_id,
                        "id": lab_id,
                        "courseName": course['course_name'],
                        "labNumber": lab['lab_number'],
                        "labTitle": lab['lab_title'],
                        "type": "lab",
                        "createdAt": timestamp,
                        "updatedAt": timestamp
                    }
                    
                    modules_table.put_item(Item=lab_item)
                    lab_count += 1
                except Exception as e:
                    print(f"실습 저장 중 오류 발생: {str(e)}")
                    failure_count += 1
            
            print(f"  - {module_count}개 모듈, {lab_count}개 실습 정보 저장 완료")
            
        except Exception as e:
            print(f"과정 {course['course_name']} 저장 중 오류 발생: {str(e)}")
            failure_count += 1
    
    print(f"\n저장 완료: 성공 {success_count}개 과정, 실패 {failure_count}건")
    return True

def save_courses_to_dynamodb_batch(json_file_path, region_name='ap-northeast-2', batch_size=25):
    """
    JSON 파일에서 과정 정보를 읽어 배치 방식으로 DynamoDB에 저장하는 함수
    
    Args:
        json_file_path (str): 과정 정보가 담긴 JSON 파일 경로
        region_name (str): AWS 리전 이름
        batch_size (int): 배치 크기 (최대 25)
    """
    # 데이터 로드
    try:
        with open(json_file_path, 'r', encoding='utf-8') as f:
            course_data = json.load(f)
        print(f"총 {len(course_data)}개 과정 정보를 로드했습니다.")
    except Exception as e:
        print(f"JSON 파일 로드 중 오류 발생: {str(e)}")
        return False
    
    # AWS 세션 및 DynamoDB 클라이언트 생성
    try:
        session = boto3.Session(region_name=region_name)
        dynamodb = session.resource('dynamodb')
        client = session.client('dynamodb')
        
        # 테이블 참조 (배치 작업 시는 테이블명만 사용)
        catalog_table_name = 'Tnc-CourseCatalog'
        modules_table_name = 'Tnc-CourseCatalog-Modules'
        
        print("AWS DynamoDB에 연결했습니다.")
    except Exception as e:
        print(f"AWS 연결 중 오류 발생: {str(e)}")
        return False
    
    # 저장 전 데이터 미리보기
    sample_course = course_data[0] if course_data else None
    if sample_course:
        print("\n=== 저장할 데이터 미리보기 ===")
        print(f"과정명: {sample_course['course_name']}")
        print(f"레벨: {sample_course['level']}")
        print(f"소요 시간: {sample_course['duration']}")
        print(f"모듈 수: {len(sample_course.get('modules', []))}")
        print(f"실습 수: {len(sample_course.get('labs', []))}")
    
    # 사용자 확인
    confirmation = input("\n위 정보를 DynamoDB에 저장하시겠습니까? (y/n): ")
    if confirmation.lower() != 'y':
        print("저장이 취소되었습니다.")
        return False
    
    # 타임스탬프 및 성공/실패 카운터
    timestamp = datetime.now().isoformat()
    success_count = 0
    failure_count = 0
    
    # 모든 항목을 배치 요청으로 구성
    catalog_batch_items = []
    modules_batch_items = []
    
    print("\n=== 배치 작업 준비 중 ===")
    for course in course_data:
        try:
            # 과정 ID 생성 (고유 식별자)
            course_id = str(uuid.uuid4())
            
            # 과정 정보 항목 생성 (DynamoDB 형식으로 변환)
            catalog_item = {
                'PutRequest': {
                    'Item': {
                        'courseId': {'S': course_id},
                        'courseName': {'S': course['course_name']},
                        'level': {'S': course['level'] if course['level'] else ''},
                        'duration': {'S': course['duration'] if course['duration'] else ''},
                        'deliveryMethod': {'S': course.get('delivery_method', '')},
                        'description': {'S': course.get('description', '')},
                        'objectives': {'S': course.get('objectives', '')},
                        'audience': {'S': course.get('audience', '')},
                        'prerequisites': {'S': course.get('prerequisites', '')},
                        'outline': {'S': course.get('outline', '')},
                        'createdAt': {'S': timestamp},
                        'updatedAt': {'S': timestamp}
                    }
                }
            }
            
            catalog_batch_items.append(catalog_item)
            
            # 모듈 정보 배치 항목 생성
            for idx, module in enumerate(course.get('modules', [])):
                module_id = f"{course_id}#module#{idx+1}"
                module_item = {
                    'PutRequest': {
                        'Item': {
                            'courseId': {'S': course_id},
                            'id': {'S': module_id},
                            'courseName': {'S': course['course_name']},
                            'moduleNumber': {'S': str(module['module_number'])},
                            'moduleTitle': {'S': module['module_title']},
                            'type': {'S': 'module'},
                            'createdAt': {'S': timestamp},
                            'updatedAt': {'S': timestamp}
                        }
                    }
                }
                modules_batch_items.append(module_item)
            
            # 실습 정보 배치 항목 생성
            for idx, lab in enumerate(course.get('labs', [])):
                lab_id = f"{course_id}#lab#{idx+1}"
                lab_item = {
                    'PutRequest': {
                        'Item': {
                            'courseId': {'S': course_id},
                            'id': {'S': lab_id},
                            'courseName': {'S': course['course_name']},
                            'labNumber': {'S': str(lab['lab_number'])},
                            'labTitle': {'S': lab['lab_title']},
                            'type': {'S': 'lab'},
                            'createdAt': {'S': timestamp},
                            'updatedAt': {'S': timestamp}
                        }
                    }
                }
                modules_batch_items.append(lab_item)
                
            success_count += 1
                
        except Exception as e:
            print(f"과정 {course['course_name']} 준비 중 오류 발생: {str(e)}")
            failure_count += 1
    
    print(f"배치 작업 준비 완료: {success_count}개 과정, {len(modules_batch_items)}개 모듈/실습 항목")
    
    # 배치 저장 작업 수행
    print("\n=== 배치 저장 작업 시작 ===")
    
    # 과정 정보 배치 저장
    catalog_batches = [catalog_batch_items[i:i+batch_size] for i in range(0, len(catalog_batch_items), batch_size)]
    for i, batch in enumerate(catalog_batches):
        try:
            response = client.batch_write_item(
                RequestItems={
                    catalog_table_name: batch
                }
            )
            print(f"과정 정보 배치 {i+1}/{len(catalog_batches)} 저장 완료 ({len(batch)}개 항목)")
            
            # 처리되지 않은 항목이 있는 경우 재시도
            unprocessed = response.get('UnprocessedItems', {}).get(catalog_table_name, [])
            retry_count = 0
            while unprocessed and retry_count < 3:
                retry_count += 1
                time.sleep(2 ** retry_count)  # 지수 백오프
                retry_response = client.batch_write_item(
                    RequestItems={
                        catalog_table_name: unprocessed
                    }
                )
                unprocessed = retry_response.get('UnprocessedItems', {}).get(catalog_table_name, [])
                print(f"  재시도 {retry_count}: {len(unprocessed)}개 항목 남음")
            
            if unprocessed:
                print(f"  경고: {len(unprocessed)}개 항목이 처리되지 않았습니다.")
                
        except Exception as e:
            print(f"과정 정보 배치 {i+1} 저장 중 오류 발생: {str(e)}")
    
    # 모듈/실습 정보 배치 저장
    modules_batches = [modules_batch_items[i:i+batch_size] for i in range(0, len(modules_batch_items), batch_size)]
    for i, batch in enumerate(modules_batches):
        try:
            response = client.batch_write_item(
                RequestItems={
                    modules_table_name: batch
                }
            )
            print(f"모듈/실습 정보 배치 {i+1}/{len(modules_batches)} 저장 완료 ({len(batch)}개 항목)")
            
            # 처리되지 않은 항목이 있는 경우 재시도
            unprocessed = response.get('UnprocessedItems', {}).get(modules_table_name, [])
            retry_count = 0
            while unprocessed and retry_count < 3:
                retry_count += 1
                time.sleep(2 ** retry_count)  # 지수 백오프
                retry_response = client.batch_write_item(
                    RequestItems={
                        modules_table_name: unprocessed
                    }
                )
                unprocessed = retry_response.get('UnprocessedItems', {}).get(modules_table_name, [])
                print(f"  재시도 {retry_count}: {len(unprocessed)}개 항목 남음")
            
            if unprocessed:
                print(f"  경고: {len(unprocessed)}개 항목이 처리되지 않았습니다.")
                
        except Exception as e:
            print(f"모듈/실습 정보 배치 {i+1} 저장 중 오류 발생: {str(e)}")
    
    print(f"\n저장 작업 완료: {len(catalog_batch_items)}개 과정, {len(modules_batch_items)}개 모듈/실습 항목 처리됨")
    return True

def verify_dynamodb_data(course_names, region_name='ap-northeast-2'):
    """
    저장된 데이터를 확인하는 함수
    
    Args:
        course_names (list): 확인할 과정명 목록
        region_name (str): AWS 리전 이름
    """
    try:
        session = boto3.Session(region_name=region_name)
        dynamodb = session.resource('dynamodb')
        
        catalog_table = dynamodb.Table('Tnc-CourseCatalog')
        modules_table = dynamodb.Table('Tnc-CourseCatalog-Modules')
        
        print("\n=== DynamoDB 저장 데이터 확인 ===")
        
        # 샘플 과정 확인 (최대 3개)
        for i, course_name in enumerate(course_names[:3]):
            # 과정 정보 조회
            response = catalog_table.query(
                IndexName='courseNameIndex',  # courseName을 인덱스로 설정한 경우
                KeyConditionExpression=Key('courseName').eq(course_name)
            )
            
            items = response.get('Items', [])
            if items:
                print(f"\n과정: {course_name}")
                course_item = items[0]
                print(f"  ID: {course_item['courseId']}")
                print(f"  레벨: {course_item.get('level', 'N/A')}")
                print(f"  소요 시간: {course_item.get('duration', 'N/A')}")
                
                # 과정에 속한 모듈/실습 조회
                modules_response = modules_table.query(
                    KeyConditionExpression=Key('courseId').eq(course_item['courseId'])
                )
                
                modules = [item for item in modules_response.get('Items', []) if item.get('type') == 'module']
                labs = [item for item in modules_response.get('Items', []) if item.get('type') == 'lab']
                
                print(f"  모듈 수: {len(modules)}")
                print(f"  실습 수: {len(labs)}")
                
                # 첫 번째 모듈과 실습 샘플 출력
                if modules:
                    print(f"  첫 번째 모듈: {modules[0].get('moduleNumber', '')} - {modules[0].get('moduleTitle', '')}")
                if labs:
                    print(f"  첫 번째 실습: {labs[0].get('labNumber', '')} - {labs[0].get('labTitle', '')}")
            else:
                print(f"\n과정 '{course_name}'을 찾을 수 없습니다.")
        
        # 전체 항목 수 조회
        catalog_scan = catalog_table.scan(Select='COUNT')
        modules_scan = modules_table.scan(Select='COUNT')
        
        print(f"\n총 저장된 과정 수: {catalog_scan['Count']}")
        print(f"총 저장된 모듈/실습 수: {modules_scan['Count']}")
        
    except Exception as e:
        print(f"데이터 확인 중 오류 발생: {str(e)}")

def extract_comprehensive_course_info(file_path):
    """
    .docx 파일에서 과정 정보, 설명, 목표, 대상, 권장 사항, 개요, 모듈, 실습명을 추출하는 함수
    (Digital Classroom 과정 제외)
    
    Args:
        file_path (str): .docx 파일 경로
        
    Returns:
        dict: 추출된 과정 정보를 담은 딕셔너리
    """
    print(f"문서 제목: {file_path}")
    
    # 문서 열기
    doc = docx.Document(file_path)
    print(f"문서를 불러왔습니다.")
    
    # 기본 정보 출력
    print(f"문단 수: {len(doc.paragraphs)}")
    print(f"표 수: {len(doc.tables)}")
    
    # 문서 요소 순차적으로 추출 (문단과 표의 실제 순서 파악)
    doc_elements = []
    for para_idx, para in enumerate(doc.paragraphs):
        doc_elements.append(("para", para_idx, para))
    for table_idx, table in enumerate(doc.tables):
        doc_elements.append(("table", table_idx, table))
    
    # 실제 문서 순서대로 요소 정렬
    try:
        all_elements_xml = doc._element.xpath('.//w:p | .//w:tbl')
        doc_elements.sort(key=lambda x: all_elements_xml.index(x[2]._element if x[0] == "para" else x[2]._element))
    except Exception as e:
        print(f"요소 정렬 중 오류: {e}")
        print("문서 요소를 원래 순서대로 유지합니다.")
    
    # 결과를 저장할 딕셔너리
    result = {
        "courses": [],
        "course_details": {},
        "course_modules": {},
        "course_labs": {},
        "course_outlines": {}
    }
    
    # 1. 2줄짜리 표 찾기 및 해당 표 바로 직전 문단에서 과정명 추출
    print("\n### 2줄짜리 표와 직전 과정명 추출 중 ###")
    
    courses = []
    for i, (elem_type, elem_idx, elem) in enumerate(doc_elements):
        if elem_type == "table" and len(elem.rows) == 2:
            table = elem
            table_idx = elem_idx
            # print(f"\n2줄짜리 표 발견 (표 #{table_idx+1})")
            
            # 과정명 추출: 표 바로 직전 문단 확인
            course_name = "과정명 추출 실패"
            for j in range(i-1, max(0, i-5), -1):  # 최대 4개 이전 요소 확인
                if j >= 0 and doc_elements[j][0] == "para":
                    prev_para = doc_elements[j][2]
                    text = prev_para.text.strip()
                    
                    # 빈 문단 건너뛰기
                    if not text:
                        continue
                        
                    # AWS 과정명 패턴 확인
                    if is_aws_course_name(text):
                        course_name = text
                        print(f"  직전 문단에서 과정명 발견: '{course_name}'")
                        break
            
            # Digital Classroom 과정 제외
            if "Digital Classroom" in course_name:
                print(f"  Digital Classroom 과정 제외: '{course_name}'")
                continue
                
            try:
                # 표 데이터 추출
                headers = []
                values = []
                
                # 첫 번째 행 (헤더)
                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip())
                
                # 두 번째 행 (값)
                for cell in table.rows[1].cells:
                    values.append(cell.text.strip())
                
                # 레벨과 소요 시간 정보 찾기
                level = None
                duration = None
                delivery_method = None
                
                # 헤더와 값을 순회하며 찾기
                for k, header in enumerate(headers):
                    if k < len(values):  # 인덱스 범위 확인
                        header_lower = header.lower()
                        if "레벨" in header_lower or "수준" in header_lower:
                            level = values[k]
                        elif "소요 시간" in header_lower or "기간" in header_lower or "duration" in header_lower:
                            duration = values[k]
                        elif "제공 방법" in header_lower or "제공 방식" in header_lower:
                            delivery_method = values[k]
                
                # 결과 저장
                course_info = {
                    "table_index": table_idx,
                    "course_name": course_name,
                    "level": level,
                    "duration": duration,
                    "delivery_method": delivery_method
                }
                
                courses.append(course_info)
                #print(f"  추출 정보: 과정명='{course_name}', 레벨='{level}', 소요 시간='{duration}'")
                
            except Exception as e:
                print(f"표 {table_idx+1} 처리 중 오류 발생: {str(e)}")
    
    # 고유한 과정명 목록 생성
    unique_course_names = []
    for course in courses:
        if course["course_name"] not in unique_course_names and course["course_name"] != "과정명 추출 실패":
            unique_course_names.append(course["course_name"])
    
    result["courses"] = courses
    
    # 2. 과정 세부 정보 추출 (설명, 목표, 대상, 권장 사항)
    print("\n### 과정 세부 정보 추출 중 ###")
    course_details = extract_course_details(doc, unique_course_names)
    result["course_details"] = course_details
    
    # 3. 모듈명 추출
    print("\n### 모듈명 추출 중 ###")
    course_modules = extract_modules(doc, exclude_digital_classroom=True, course_names=unique_course_names)
    result["course_modules"] = course_modules
    
    # 4. 실습명 추출
    print("\n### 실습명 추출 중 ###")
    course_labs = extract_labs(doc, exclude_digital_classroom=True, course_names=unique_course_names)
    result["course_labs"] = course_labs
    
    # 5. 과정 개요 추출
    print("\n### 과정 개요 추출 중 ###")
    course_outlines = extract_course_outlines(doc, unique_course_names)
    result["course_outlines"] = course_outlines
    
    return result

def is_aws_course_name(text):
    """
    텍스트가 AWS 과정명인지 확인
    
    Args:
        text: 확인할 텍스트
        
    Returns:
        bool: AWS 과정명 여부
    """
    # 과정명으로 볼 수 있는 조건
    if not text or len(text) < 5 or len(text) > 100:
        return False
    
    # 1. 일반적인 제외 패턴
    exclude_patterns = [
        r'^모듈',
        r'^실습',
        r'^•',
        r'^-',
        r'^참고:',
        r'^주의:',
        r'^과정 설명',
        r'^Digital Classroom',  # Digital Classroom으로 시작하는 경우
    ]
    
    for pattern in exclude_patterns:
        if re.match(pattern, text):
            return False
    
    # 2. AWS 과정명 패턴 확인 (우선순위 순)
    aws_patterns = [
        r'^AWS .+',                 # AWS로 시작하는 과정명
        r'^Amazon .+',              # Amazon으로 시작하는 과정명
        r'.+ on AWS\$',              # on AWS로 끝나는 과정명
        r'.*AWS.*',                 # AWS가 포함된 과정명
        r'.*Amazon.*',              # Amazon이 포함된 과정명
        r'.*Cloud.*',               # Cloud가 포함된 과정명
    ]
    
    for pattern in aws_patterns:
        if re.match(pattern, text):
            return True
    
    return False

def extract_course_details(doc, course_names):
    """
    문서에서 과정 설명, 목표, 대상, 권장 사항 추출
    
    Args:
        doc: 문서 객체
        course_names: 과정명 목록
        
    Returns:
        dict: 과정별 상세 정보
    """
    details = {}
    
    # 과정 상세 정보 섹션 헤더 패턴
    section_patterns = {
        "description": [r'과정\s*설명', r'개요'],
        "objectives": [r'과정\s*목표', r'교육\s*목표', r'이\s?과정\s?에서\s?배우게\s?될\s?내용'],
        "audience": [r'수강\s*대상', r'교육\s*대상', r'이\s?과정의\s?대상은\s?다음과\s?같습니다'],
        "prerequisites": [r'수강\s*전\s*권장\s*사항', r'선수\s*과목', r'이\s?과정을\s?수강하려면\s?다음\s?조건을\s?갖추는\s?것이\s?좋습니다']
    }
    
    # 현재 처리 중인 과정과 섹션
    current_course = None
    current_section = None
    current_content = []
    
    # 문단 순회
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 과정명 확인
        matched_course = None
        if is_aws_course_name(text):
            for course_name in course_names:
                if text == course_name or (len(course_name) > 10 and (course_name in text or text in course_name)):
                    matched_course = course_name
                    break
        
        # 새 과정 발견
        if matched_course:
            # 이전 과정의 섹션 저장
            if current_course and current_section and current_content:
                if current_course not in details:
                    details[current_course] = {}
                details[current_course][current_section] = '\n'.join(current_content)
            
            current_course = matched_course
            current_section = None
            current_content = []
            continue
        
        # 현재 과정이 설정된 상태에서 섹션 헤더 확인
        if current_course:
            matched_section = None
            for section, patterns in section_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, text, re.IGNORECASE):
                        matched_section = section
                        break
                if matched_section:
                    break
            
            # 새 섹션 헤더 발견
            if matched_section:
                # 이전 섹션 내용 저장
                if current_section and current_content:
                    if current_course not in details:
                        details[current_course] = {}
                    details[current_course][current_section] = '\n'.join(current_content)
                
                current_section = matched_section
                current_content = []
                continue
            
            # 현재 섹션에 내용 추가
            if current_section:
                # 불릿 포인트 형태 정리
                text = re.sub(r'^\s*[•\-]\s*', '• ', text)  # 불릿 포인트 통일
                current_content.append(text)
    
    # 마지막 과정/섹션 저장
    if current_course and current_section and current_content:
        if current_course not in details:
            details[current_course] = {}
        details[current_course][current_section] = '\n'.join(current_content)
    
    return details

def extract_course_outlines(doc, course_names):
    """
    문서에서 과정 개요 추출
    
    Args:
        doc: 문서 객체
        course_names: 과정명 목록
        
    Returns:
        dict: 과정별 개요 정보
    """
    outlines = {}
    
    # 과정 개요 시작/끝 패턴
    outline_start_patterns = [r'과정\s*개요', r'교육\s*과정\s*개요']
    outline_end_patterns = [r'과정\s*마무리', r'리소스', r'사후\s*평가']
    
    # 현재 과정과 개요 수집 상태
    current_course = None
    collecting_outline = False
    current_outline = []
    
    # 문단 순회
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 과정명 확인
        matched_course = None
        if is_aws_course_name(text):
            for course_name in course_names:
                if text == course_name or (len(course_name) > 10 and (course_name in text or text in course_name)):
                    matched_course = course_name
                    break
        
        # 새 과정 발견
        if matched_course:
            # 이전 과정의 개요 저장
            if current_course and collecting_outline and current_outline:
                outlines[current_course] = '\n'.join(current_outline)
            
            current_course = matched_course
            collecting_outline = False
            current_outline = []
            continue
        
        # 개요 섹션 시작 확인
        if current_course and not collecting_outline:
            for pattern in outline_start_patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    collecting_outline = True
                    break
            
            if collecting_outline:
                continue  # 시작 헤더는 개요에 포함하지 않음
        
        # 개요 섹션 끝 확인
        if current_course and collecting_outline:
            for pattern in outline_end_patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    collecting_outline = False
                    break
            
            if not collecting_outline:
                continue  # 종료 헤더는 개요에 포함하지 않음
        
        # 개요 내용 수집
        if current_course and collecting_outline:
            # 모듈 헤더 형식 표준화
            text = re.sub(r'^모듈\s*(\d+)[:\s-]', r'모듈 \1: ', text)
            current_outline.append(text)
    
    # 마지막 과정의 개요 저장
    if current_course and collecting_outline and current_outline:
        outlines[current_course] = '\n'.join(current_outline)
    
    return outlines

def extract_modules(doc, exclude_digital_classroom=True, course_names=None):
    """
    문서에서 모듈 정보 추출
    
    Args:
        doc: 문서 객체
        exclude_digital_classroom: Digital Classroom 과정 제외 여부
        course_names: 이미 식별된 과정명 목록
        
    Returns:
        dict: 과정별 모듈 정보
    """
    modules = {}
    current_course = None
    current_modules = []
    
    # 과정명 목록이 제공된 경우 우선 확인
    course_list = []
    if course_names:
        course_list = [name for name in course_names if name != "과정명 추출 실패"]
    
    # 모듈 패턴
    module_pattern = re.compile(r'모듈\s*(\d+)[:\s-]\s*(.+)', re.IGNORECASE)
    module_pattern_eng = re.compile(r'Module\s*(\d+)[:\s-]\s*(.+)', re.IGNORECASE)
    module_day_pattern = re.compile(r'(\d+)일\s*차\s*[:：]?\s*(.+)', re.IGNORECASE)  # "1일 차: 내용" 형태
    
    # 현재 과정명 찾기
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 과정명으로 보이는 문단 확인
        if is_aws_course_name(text):
            # Digital Classroom 과정 제외
            if exclude_digital_classroom and "Digital Classroom" in text:
                current_course = None  # 현재 과정 초기화
                current_modules = []
                continue
            
            # 이미 식별된 과정명 중에 있는지 확인
            matched_course = None
            if course_list:
                for course in course_list:
                    # 정확히 일치하는 과정명 또는 유사한 과정명 확인
                    if course == text or (len(course) > 10 and (course in text or text in course)):
                        matched_course = course
                        break
            
            # 이전 과정의 모듈을 저장
            if current_course and current_modules:
                modules[current_course] = current_modules
            
            # 새 과정 시작
            current_course = matched_course if matched_course else text
            current_modules = []
            continue
        
        # 모듈 패턴 확인
        module_match = module_pattern.match(text) or module_pattern_eng.match(text) or module_day_pattern.match(text)
        if module_match and current_course:
            module_num = module_match.group(1)
            module_title = module_match.group(2).strip()
            
            # 모듈명이 너무 길거나 불릿 포인트로 시작하면 제외
            if len(module_title) < 200 and not module_title.startswith('•'):
                current_modules.append({
                    "module_number": module_num,
                    "module_title": module_title
                })
    
    # 마지막 과정의 모듈 저장
    if current_course and current_modules:
        modules[current_course] = current_modules
    
    return modules

def extract_labs(doc, exclude_digital_classroom=True, course_names=None):
    """
    문서에서 실습 정보 추출
    
    Args:
        doc: 문서 객체
        exclude_digital_classroom: Digital Classroom 과정 제외 여부
        course_names: 이미 식별된 과정명 목록
        
    Returns:
        dict: 과정별 실습 정보
    """
    labs = {}
    current_course = None
    current_labs = []
    
    # 과정명 목록이 제공된 경우 우선 확인
    course_list = []
    if course_names:
        course_list = [name for name in course_names if name != "과정명 추출 실패"]
    
    # 실습 패턴
    lab_pattern = re.compile(r'실습\s*(\d+)[:\s-]\s*(.+)', re.IGNORECASE)
    lab_pattern_eng = re.compile(r'Lab\s*(\d+)[:\s-]\s*(.+)', re.IGNORECASE)
    hands_on_pattern = re.compile(r'핸즈온\s*랩\s*(\d*)[:\s-]?\s*(.+)', re.IGNORECASE)  # "핸즈온랩: 제목" 형태
    
    # 현재 과정명 찾기
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 과정명으로 보이는 문단 확인
        if is_aws_course_name(text):
            # Digital Classroom 과정 제외
            if exclude_digital_classroom and "Digital Classroom" in text:
                current_course = None  # 현재 과정 초기화
                current_labs = []
                continue
            
            # 이미 식별된 과정명 중에 있는지 확인
            matched_course = None
            if course_list:
                for course in course_list:
                    # 정확히 일치하는 과정명 또는 유사한 과정명 확인
                    if course == text or (len(course) > 10 and (course in text or text in course)):
                        matched_course = course
                        break
            
            # 이전 과정의 실습을 저장
            if current_course and current_labs:
                labs[current_course] = current_labs
            
            # 새 과정 시작
            current_course = matched_course if matched_course else text
            current_labs = []
            continue
        
        # 실습 패턴 확인
        lab_match = lab_pattern.match(text) or lab_pattern_eng.match(text) or hands_on_pattern.match(text)
        if lab_match and current_course:
            lab_num = lab_match.group(1) if lab_match.group(1) else "N/A"
            lab_title = lab_match.group(2).strip()
            
            # 실습명이 너무 길거나 불릿 포인트로 시작하면 제외
            if len(lab_title) < 200 and not lab_title.startswith('•'):
                current_labs.append({
                    "lab_number": lab_num,
                    "lab_title": lab_title
                })
    
    # 마지막 과정의 실습 저장
    if current_course and current_labs:
        labs[current_course] = current_labs
    
    return labs

def display_comprehensive_course_info(result):
    """
    추출된 종합 과정 정보 출력
    
    Args:
        result: 추출된 종합 과정 정보
    """
    courses = result["courses"]
    course_details = result["course_details"]
    course_modules = result["course_modules"]
    course_labs = result["course_labs"]
    course_outlines = result["course_outlines"]
    
    # Digital Classroom 과정 필터링
    filtered_courses = [info for info in courses if "Digital Classroom" not in info["course_name"]]
    
    # 과정 정보 요약 DataFrame
    if filtered_courses:
        summary_data = []
        for info in filtered_courses:
            course_name = info["course_name"]
            if course_name == "과정명 추출 실패":
                continue
                
            # 과정 세부 정보 확인
            description = ""
            if course_name in course_details and "description" in course_details[course_name]:
                description = course_details[course_name]["description"]
            
            # 요약 데이터 생성
            summary_data.append({
                "과정명": course_name,
                "레벨": info["level"] if info["level"] else "",
                "소요 시간": info["duration"] if info["duration"] else "",
                "제공 방법": info["delivery_method"] if info.get("delivery_method") else "",
                "설명": description[:100] + "..." if len(description) > 100 else description
            })
        
        df_summary = pd.DataFrame(summary_data)
        print(f"\n=== 과정 정보 요약 (Digital Classroom 제외, {len(df_summary)}개 과정) ===")
        pd.set_option('display.max_colwidth', 30)  # 열 너비 제한
        print(df_summary)
        
        # CSV로 저장
        df_summary.to_csv("course_summary.csv", index=False, encoding='utf-8-sig')
        print("과정 정보 요약을 'course_summary.csv'에 저장했습니다.")
    
    # 과정 상세 정보 저장
    detailed_data = []
    for info in filtered_courses:
        course_name = info["course_name"]
        if course_name == "과정명 추출 실패":
            continue
        
        course_data = {
            "과정명": course_name,
            "레벨": info["level"] if info["level"] else "",
            "소요 시간": info["duration"] if info["duration"] else "",
            "제공 방법": info["delivery_method"] if info.get("delivery_method") else "",
            "설명": course_details.get(course_name, {}).get("description", ""),
            "과정 목표": course_details.get(course_name, {}).get("objectives", ""),
            "수강 대상": course_details.get(course_name, {}).get("audience", ""),
            "수강 전 권장 사항": course_details.get(course_name, {}).get("prerequisites", ""),
            "모듈 수": len(course_modules.get(course_name, [])),
            "실습 수": len(course_labs.get(course_name, []))
        }
        detailed_data.append(course_data)
    
    if detailed_data:
        # CSV로 저장
        pd.DataFrame(detailed_data).to_csv("course_details.csv", index=False, encoding='utf-8-sig')
        print("과정 세부 정보를 'course_details.csv'에 저장했습니다.")
    
    # 과정 개요 저장
    outline_data = []
    for course_name, outline in course_outlines.items():
        if "Digital Classroom" in course_name:
            continue
            
        outline_data.append({
            "과정명": course_name,
            "개요": outline
        })
    
    if outline_data:
        pd.DataFrame(outline_data).to_csv("course_outlines.csv", index=False, encoding='utf-8-sig')
        print("과정 개요를 'course_outlines.csv'에 저장했습니다.")
    
    # 모듈 정보 저장
    modules_data = []
    for course_name, modules in course_modules.items():
        if "Digital Classroom" in course_name:
            continue
            
        for module in modules:
            modules_data.append({
                "과정명": course_name,
                "모듈 번호": module["module_number"],
                "모듈 제목": module["module_title"]
            })
    
    if modules_data:
        pd.DataFrame(modules_data).to_csv("course_modules.csv", index=False, encoding='utf-8-sig')
        print("모듈 정보를 'course_modules.csv'에 저장했습니다.")
    
    # 실습 정보 저장
    labs_data = []
    for course_name, labs in course_labs.items():
        if "Digital Classroom" in course_name:
            continue
            
        for lab in labs:
            labs_data.append({
                "과정명": course_name,
                "실습 번호": lab["lab_number"],
                "실습 제목": lab["lab_title"]
            })
    
    if labs_data:
        pd.DataFrame(labs_data).to_csv("course_labs.csv", index=False, encoding='utf-8-sig')
        print("실습 정보를 'course_labs.csv'에 저장했습니다.")
    
    # 특정 과정의 모든 정보 샘플 출력
    if filtered_courses:
        sample_course = filtered_courses[0]["course_name"]
        print(f"\n=== 과정 샘플 정보: {sample_course} ===")
        
        # 기본 정보
        print("레벨:", filtered_courses[0]["level"] if filtered_courses[0]["level"] else "정보 없음")
        print("소요 시간:", filtered_courses[0]["duration"] if filtered_courses[0]["duration"] else "정보 없음")
        
        # 설명
        if sample_course in course_details and "description" in course_details[sample_course]:
            print("\n[설명]")
            print(course_details[sample_course]["description"][:200] + "..." if len(course_details[sample_course]["description"]) > 200 else course_details[sample_course]["description"])
        
        # 과정 목표
        if sample_course in course_details and "objectives" in course_details[sample_course]:
            print("\n[과정 목표]")
            print(course_details[sample_course]["objectives"][:200] + "..." if len(course_details[sample_course]["objectives"]) > 200 else course_details[sample_course]["objectives"])
        
        # 수강 대상
        if sample_course in course_details and "audience" in course_details[sample_course]:
            print("\n[수강 대상]")
            print(course_details[sample_course]["audience"][:200] + "..." if len(course_details[sample_course]["audience"]) > 200 else course_details[sample_course]["audience"])
        
        # 수강 전 권장 사항
        if sample_course in course_details and "prerequisites" in course_details[sample_course]:
            print("\n[수강 전 권장 사항]")
            print(course_details[sample_course]["prerequisites"][:200] + "..." if len(course_details[sample_course]["prerequisites"]) > 200 else course_details[sample_course]["prerequisites"])
        
        # 모듈 정보
        if sample_course in course_modules and course_modules[sample_course]:
            print("\n[모듈 정보]")
            for i, module in enumerate(course_modules[sample_course][:3]):  # 처음 3개만
                print(f"  모듈 {module['module_number']}: {module['module_title']}")
            if len(course_modules[sample_course]) > 3:
                print(f"  ... 외 {len(course_modules[sample_course]) - 3}개 모듈")
        
        # 과정 개요
        if sample_course in course_outlines:
            print("\n[과정 개요]")
            outline = course_outlines[sample_course]
            print(outline[:200] + "..." if len(outline) > 200 else outline)
    
    # JSON 형식으로 모든 정보 저장
    all_course_data = []
    
    for info in filtered_courses:
        course_name = info["course_name"]
        if course_name == "과정명 추출 실패":
            continue
        
        # 과정 상세 정보 수집
        course_dict = {
            "course_name": course_name,
            "level": info["level"] if info["level"] else "",
            "duration": info["duration"] if info["duration"] else "",
            "delivery_method": info["delivery_method"] if info.get("delivery_method") else "",
            "description": course_details.get(course_name, {}).get("description", ""),
            "objectives": course_details.get(course_name, {}).get("objectives", ""),
            "audience": course_details.get(course_name, {}).get("audience", ""),
            "prerequisites": course_details.get(course_name, {}).get("prerequisites", ""),
            "outline": course_outlines.get(course_name, ""),
            "modules": course_modules.get(course_name, []),
            "labs": course_labs.get(course_name, [])
        }
        
        all_course_data.append(course_dict)
    
    # JSON 파일로 저장
    with open("all_course_data.json", "w", encoding="utf-8") as f:
        json.dump(all_course_data, f, ensure_ascii=False, indent=2)
    
    print("\n모든 과정 데이터를 'all_course_data.json'에 저장했습니다.")

def main():
    """메인 실행 함수"""
    try:
        import sys
        if len(sys.argv) > 1:
            file_path = sys.argv[1]
        else:
            file_path = 'AWS TnC_ILT_DILT.docx'  # 기본 파일 이름
        
        start_time = time.time()
        result = extract_comprehensive_course_info(file_path)
        end_time = time.time()
        
        print(f"\n처리 완료! 소요 시간: {end_time - start_time:.2f}초")
        display_comprehensive_course_info(result)
        
        # 문서 내용 전체 저장 (선택적)
        doc = docx.Document(file_path)
        with open("extracted_text.txt", "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
            print("\n문서 내용을 extracted_text.txt 파일로 저장했습니다.")
            
    except Exception as e:
        print(f"오류 발생: {str(e)}")
        import traceback
        traceback.print_exc()

    # try:
    #     # 사용할 JSON 파일 결정
    #     json_file_path = "all_course_data.json"
    #     print(f"'{json_file_path}' 파일을 DynamoDB에 저장합니다.")
        
    #     # 배치 방식으로 저장 (더 효율적)
    #     result = save_courses_to_dynamodb_batch(json_file_path)
        
    #     if result:
    #         # 저장된 데이터 확인
    #         try:
    #             with open(json_file_path, 'r', encoding='utf-8') as f:
    #                 course_data = json.load(f)
    #             course_names = [course['course_name'] for course in course_data]
    #             verify_dynamodb_data(course_names)
    #         except Exception as e:
    #             print(f"저장된 데이터 확인 중 오류 발생: {str(e)}")
        
    # except Exception as e:
    #     print(f"오류 발생: {str(e)}")
    #     import traceback
    #     traceback.print_exc()

if __name__ == "__main__":
    main()